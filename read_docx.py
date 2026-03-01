import docx
import sys

# Set output encoding to utf-8
if sys.stdout.encoding != 'utf-8':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc = docx.Document("f:/code/web/MMC/frontend/MMC_Volunteer_Registration_Form.docx")
for p in doc.paragraphs:
    if p.text.strip():
        print(p.text.strip())
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                print('[TABLE]', cell.text.strip())
