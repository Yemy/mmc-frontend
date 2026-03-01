import docx
import os

def read_docx(file_path):
    if not os.path.exists(file_path):
        return f"File not found: {file_path}"
    doc = docx.Document(file_path)
    content = []
    for p in doc.paragraphs:
        if p.text.strip():
            content.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                content.append("[TABLE] " + " | ".join(row_data))
    return "\n".join(content)

with open("f:/code/web/MMC/frontend/mmc_description.txt", "w", encoding="utf-8") as f:
    f.write(read_docx("f:/code/web/MMC/frontend/MMC description.docx"))

with open("f:/code/web/MMC/frontend/projects_implemented.txt", "w", encoding="utf-8") as f:
    f.write(read_docx("f:/code/web/MMC/frontend/projects implemented by MMC.docx"))

print("Extracted content saved to mmc_description.txt and projects_implemented.txt")
